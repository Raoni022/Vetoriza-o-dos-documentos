from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import SupabaseVectorStore
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_openai import ChatOpenAI
from supabase import create_client
import os
import docx
from io import BytesIO

app = FastAPI()

# Variáveis de ambiente e API keys
SUPABASE_URL = "https://kjwdegxainpuekfwajre.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Imtqd2RlZ3hhaW5wdWVrZndhanJlIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc0NjYxOTY0MSwiZXhwIjoyMDYyMTk1NjQxfQ.8tOq83Bp68oSXslYWtrnYrLAJjTZZ0esH2A5uioDqOk"
OPENAI_API_KEY = "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"

os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

vector_store = SupabaseVectorStore(
    client=supabase,
    embedding=OpenAIEmbeddings(model="text-embedding-3-small"),
    table_name="documentos_embeddings",
    query_name="match_documents_embeddings"
)

retriever = vector_store.as_retriever(search_kwargs={"k": 5})

qa_chain = RetrievalQA.from_chain_type(
    llm=ChatOpenAI(model="gpt-4", temperature=0),
    retriever=retriever,
    return_source_documents=True
)

class Pergunta(BaseModel):
    pergunta: str

@app.post("/responder")
async def responder(p: Pergunta):
    try:
        resultado = qa_chain.invoke({"query": p.pergunta})
        return {
            "resposta": resultado["result"],
            "fontes": [doc.metadata for doc in resultado["source_documents"]]
        }
    except Exception as e:
        return {
            "erro": str(e),
            "mensagem": "Verifique se a função match_documents_embeddings existe e se os embeddings estão corretamente inseridos."
        }

@app.post("/melhorar_documento")
async def melhorar_documento(file: UploadFile = File(...)):
    try:
        doc_bytes = await file.read()
        doc = docx.Document(BytesIO(doc_bytes))
        texto_original = "\n".join([p.text for p in doc.paragraphs])

        prompt = f"Melhore o texto a seguir mantendo estilo e estrutura: \n\n{texto_original}"
        resposta = ChatOpenAI(model="gpt-4", temperature=0).invoke(prompt)

        linhas_novas = resposta.content.split("\n")
        for i, p in enumerate(doc.paragraphs):
            if i < len(linhas_novas):
                p.text = linhas_novas[i]

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        return {
            "mensagem": "Documento melhorado com sucesso.",
            "arquivo_base64": output_stream.getvalue().hex()
        }
    except Exception as e:
        return {"erro": str(e)}

